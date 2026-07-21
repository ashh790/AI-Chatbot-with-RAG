"""Extract text from uploaded documents and index them for RAG.

Supported: .pdf .docx .txt .md .csv .json .log

SECURITY NOTE
-------------
Filenames arrive from the browser and are attacker-controlled. A name like
"../../.env" or "C:\\Windows\\system32\\x" would otherwise let an upload write
outside docs/. safe_filename() strips every path component and rejects anything
that still escapes the target directory after resolution.
"""

from __future__ import annotations

import json
import logging
import re
import unicodedata
from pathlib import Path

from .config import Config
from .utils import chunk_text, vector_store

logger = logging.getLogger(__name__)

PDF_EXT = {".pdf"}
DOCX_EXT = {".docx"}
TEXT_EXT = {".txt", ".md", ".markdown", ".csv", ".json", ".log"}
ALLOWED_EXT = PDF_EXT | DOCX_EXT | TEXT_EXT

MAX_UPLOAD_BYTES = 25 * 1024 * 1024  # 25 MB


class DocumentError(ValueError):
    """Anything wrong with an uploaded document."""


# ---------------------------------------------------------------------------
# Filename safety
# ---------------------------------------------------------------------------

_UNSAFE = re.compile(r"[^A-Za-z0-9._ \-()]+")


def safe_filename(raw: str) -> str:
    """Reduce a browser-supplied filename to a bare, safe basename."""
    if not raw or not raw.strip():
        raise DocumentError("filename is empty")

    # Normalise unicode so look-alike characters can't smuggle separators.
    name = unicodedata.normalize("NFKC", raw).strip()

    # Take the basename under BOTH separator conventions -- a Windows client
    # can send backslashes that posixpath would happily treat as content.
    name = name.replace("\\", "/").split("/")[-1]

    # Drop anything that isn't clearly safe, collapse runs.
    name = _UNSAFE.sub("_", name)
    name = re.sub(r"_{2,}", "_", name).strip(". ")

    if not name or name in {".", ".."}:
        raise DocumentError(f"filename {raw!r} is not usable")
    if len(name) > 120:
        stem, dot, ext = name.rpartition(".")
        name = (stem[:110] + dot + ext) if dot else name[:120]
    return name


def resolve_within(directory: Path, filename: str) -> Path:
    """Join and confirm the result really is inside `directory`."""
    directory = directory.resolve()
    target = (directory / safe_filename(filename)).resolve()
    if directory not in target.parents and target.parent != directory:
        raise DocumentError("resolved path escapes the documents directory")
    return target


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

MAX_PDF_PAGES = 800
MAX_EXTRACTED_CHARS = 5_000_000


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF without letting a bad file kill the process.

    Design-heavy or malformed PDFs can make pypdf allocate enormous amounts of
    memory or recurse deeply. An unhandled MemoryError/RecursionError inside a
    request takes the whole dev server down, which the browser reports as
    ERR_CONNECTION_RESET -- so both are caught per page and overall.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentError("pypdf is not installed. Run: pip install pypdf") from exc

    try:
        reader = PdfReader(str(path), strict=False)
    except MemoryError as exc:
        raise DocumentError("PDF is too complex to parse (out of memory)") from exc
    except RecursionError as exc:
        raise DocumentError("PDF structure is malformed (too deeply nested)") from exc
    except Exception as exc:
        raise DocumentError(f"could not open PDF: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            if reader.decrypt("") == 0:
                raise DocumentError("PDF is password-protected")
        except DocumentError:
            raise
        except Exception as exc:
            raise DocumentError(f"PDF is encrypted and could not be opened: {exc}") from exc

    try:
        page_count = len(reader.pages)
    except Exception as exc:
        raise DocumentError(f"could not read PDF page list: {exc}") from exc

    if page_count > MAX_PDF_PAGES:
        raise DocumentError(
            f"PDF has {page_count} pages; the limit is {MAX_PDF_PAGES}. "
            "Split it into smaller files."
        )

    pages, total, failed = [], 0, 0
    for i in range(page_count):
        try:
            text = reader.pages[i].extract_text() or ""
        except (MemoryError, RecursionError) as exc:
            logger.warning("page %d aborted (%s)", i + 1, type(exc).__name__)
            failed += 1
            continue
        except Exception as exc:
            logger.warning("page %d unreadable: %s", i + 1, exc)
            failed += 1
            continue

        pages.append(text)
        total += len(text)
        if total > MAX_EXTRACTED_CHARS:
            logger.warning("stopping at page %d -- extraction limit reached", i + 1)
            break

    if failed:
        logger.warning("%d of %d page(s) could not be read in %s", failed, page_count, path.name)
    if failed == page_count:
        raise DocumentError(
            f"none of the {page_count} pages could be read. The PDF may be "
            "corrupt, or a scan containing images rather than text."
        )
    return "\n\n".join(pages)


def extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise DocumentError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocumentError(f"could not open DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables carry a lot of the content in real-world documents.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            raw = path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            raise DocumentError(f"could not read file: {exc}") from exc

        if path.suffix.lower() == ".json":
            try:
                return json.dumps(json.loads(raw), indent=2)
            except json.JSONDecodeError:
                return raw
        return raw
    raise DocumentError("could not decode file with any known text encoding")


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in PDF_EXT:
        return extract_pdf(path)
    if ext in DOCX_EXT:
        return extract_docx(path)
    if ext in TEXT_EXT:
        return extract_text_file(path)
    raise DocumentError(f"unsupported file type '{ext}'")


# ---------------------------------------------------------------------------
# Ingest
# ---------------------------------------------------------------------------

def ingest_file(path: Path, source_name: str | None = None) -> dict:
    """Extract, chunk and index one file. Returns a summary dict."""
    source = source_name or path.name
    text = extract_text(path)

    if not text.strip():
        raise DocumentError(
            "no extractable text found. If this is a scanned PDF, it contains "
            "images rather than text -- run OCR on it first."
        )

    chunks = chunk_text(text)
    if not chunks:
        raise DocumentError("document produced no usable chunks")

    if not vector_store.available:
        raise DocumentError(f"vector store unavailable: {vector_store.error}")

    ids = [f"{source}::chunk-{i}" for i in range(len(chunks))]
    metas = [
        {"source": source, "chunk": i, "suffix": path.suffix.lower()}
        for i in range(len(chunks))
    ]
    stored = vector_store.add_documents(chunks, ids, metas)

    return {
        "source": source,
        "chunks": stored,
        "characters": len(text),
        "bytes": path.stat().st_size if path.exists() else 0,
    }


def save_and_ingest(file_storage, docs_dir: Path | None = None) -> dict:
    """Handle a Werkzeug FileStorage: validate, save under docs/, index it."""
    docs_dir = Path(docs_dir or Config.DOCS_DIR)
    docs_dir.mkdir(parents=True, exist_ok=True)

    raw_name = file_storage.filename or ""
    name = safe_filename(raw_name)
    ext = Path(name).suffix.lower()

    if ext not in ALLOWED_EXT:
        raise DocumentError(
            f"'{ext or 'no extension'}' is not supported. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXT))}"
        )

    # Re-uploading the same filename REPLACES it. Creating report (2).pdf,
    # report (3).pdf ... on every retry silently multiplies the knowledge base
    # and makes the model cite whichever copy it happens to retrieve.
    target = resolve_within(docs_dir, name)
    replaced = target.exists()
    if replaced:
        vector_store.delete_by_source(target.name)
        logger.info("Replacing existing document %s", target.name)

    file_storage.save(str(target))

    size = target.stat().st_size
    if size == 0:
        target.unlink(missing_ok=True)
        raise DocumentError("file is empty")
    if size > MAX_UPLOAD_BYTES:
        target.unlink(missing_ok=True)
        MB = 1024 * 1024
        raise DocumentError(
            f"file is {size / MB:.1f} MB; limit is {MAX_UPLOAD_BYTES // MB} MB"
        )

    try:
        result = ingest_file(target, source_name=target.name)
    except Exception:
        # Don't leave an unindexed orphan behind if extraction failed.
        target.unlink(missing_ok=True)
        raise

    result["saved_as"] = target.name
    result["replaced"] = replaced
    return result


def delete_document(source: str, docs_dir: Path | None = None) -> dict:
    """Remove a document's chunks from the index and its file from disk."""
    docs_dir = Path(docs_dir or Config.DOCS_DIR)
    removed_chunks = vector_store.delete_by_source(source)

    file_removed = False
    try:
        target = resolve_within(docs_dir, source)
        if target.exists():
            target.unlink()
            file_removed = True
    except DocumentError:
        pass  # a source that isn't a real filename simply has no file

    return {
        "source": source,
        "chunks_removed": removed_chunks,
        "file_removed": file_removed,
    }
